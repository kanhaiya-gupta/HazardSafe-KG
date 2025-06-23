"""
Document Processing Module for RAG System

This module handles the extraction and processing of documents (PDFs, Word docs, etc.)
into structured JSON format for the RAG system.
"""

import logging
from typing import Dict, List, Any, Optional, Union
from pathlib import Path
import json
import hashlib
import uuid
from datetime import datetime
import re

# PDF processing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word document processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Text processing
try:
    import nltk
    from nltk.tokenize import sent_tokenize, word_tokenize
    from nltk.corpus import stopwords
    NLTK_AVAILABLE = True
except ImportError:
    NLTK_AVAILABLE = False

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Processes documents and extracts structured information for RAG."""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._extract_from_pdf,
            '.docx': self._extract_from_docx,
            '.txt': self._extract_from_txt,
            '.json': self._extract_from_json
        }
        
        # Initialize NLTK if available
        if NLTK_AVAILABLE:
            try:
                nltk.download('punkt', quiet=True)
                nltk.download('stopwords', quiet=True)
                nltk.download('averaged_perceptron_tagger', quiet=True)
            except:
                pass
    
    def process_document(self, file_path: str, doc_type: str = "general") -> Dict[str, Any]:
        """
        Process a document and extract structured information.
        
        Args:
            file_path: Path to the document file
            doc_type: Type of document (safety, engineering, regulatory, etc.)
            
        Returns:
            Dictionary containing extracted document information
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                return {
                    "success": False,
                    "error": f"File not found: {file_path}"
                }
            
            # Check if format is supported
            file_ext = file_path.suffix.lower()
            if file_ext not in self.supported_formats:
                return {
                    "success": False,
                    "error": f"Unsupported file format: {file_ext}"
                }
            
            # Extract content based on file type
            extraction_func = self.supported_formats[file_ext]
            extracted_content = extraction_func(file_path)
            
            if not extracted_content["success"]:
                return extracted_content
            
            # Process and structure the content
            processed_doc = self._structure_document(
                extracted_content["content"],
                file_path,
                doc_type
            )
            
            return {
                "success": True,
                "document": processed_doc
            }
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return {
                "success": False,
                "error": f"Processing error: {str(e)}"
            }
    
    def _extract_from_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text and metadata from PDF files."""
        if not PDF_AVAILABLE:
            return {
                "success": False,
                "error": "PDF processing not available. Install PyPDF2 and pdfplumber."
            }
        
        try:
            content = {
                "text": "",
                "metadata": {},
                "pages": []
            }
            
            # Extract text using pdfplumber (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                content["metadata"] = {
                    "pages": len(pdf.pages),
                    "title": pdf.metadata.get("Title", ""),
                    "author": pdf.metadata.get("Author", ""),
                    "subject": pdf.metadata.get("Subject", ""),
                    "creator": pdf.metadata.get("Creator", ""),
                    "producer": pdf.metadata.get("Producer", ""),
                    "creation_date": pdf.metadata.get("CreationDate", ""),
                    "modification_date": pdf.metadata.get("ModDate", "")
                }
                
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        content["text"] += f"\n\n--- Page {page_num + 1} ---\n{page_text}"
                        content["pages"].append({
                            "page_number": page_num + 1,
                            "text": page_text,
                            "words": len(page_text.split())
                        })
            
            # Fallback to PyPDF2 if pdfplumber didn't extract text
            if not content["text"].strip():
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    content["metadata"]["pages"] = len(pdf_reader.pages)
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            content["text"] += f"\n\n--- Page {page_num + 1} ---\n{page_text}"
                            content["pages"].append({
                                "page_number": page_num + 1,
                                "text": page_text,
                                "words": len(page_text.split())
                            })
            
            return {
                "success": True,
                "content": content
            }
            
        except Exception as e:
            logger.error(f"Error extracting from PDF {file_path}: {e}")
            return {
                "success": False,
                "error": f"PDF extraction error: {str(e)}"
            }
    
    def _extract_from_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text and metadata from Word documents."""
        if not DOCX_AVAILABLE:
            return {
                "success": False,
                "error": "Word document processing not available. Install python-docx."
            }
        
        try:
            doc = Document(file_path)
            
            content = {
                "text": "",
                "metadata": {
                    "paragraphs": len(doc.paragraphs),
                    "sections": len(doc.sections),
                    "tables": len(doc.tables)
                },
                "paragraphs": [],
                "tables": []
            }
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content["text"] += para.text + "\n"
                    content["paragraphs"].append({
                        "text": para.text,
                        "style": para.style.name,
                        "alignment": str(para.alignment)
                    })
            
            # Extract text from tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                    content["text"] += " | ".join(row_data) + "\n"
                content["tables"].append(table_data)
            
            return {
                "success": True,
                "content": content
            }
            
        except Exception as e:
            logger.error(f"Error extracting from Word document {file_path}: {e}")
            return {
                "success": False,
                "error": f"Word document extraction error: {str(e)}"
            }
    
    def _extract_from_txt(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            content = {
                "text": text,
                "metadata": {
                    "lines": len(text.split('\n')),
                    "words": len(text.split()),
                    "characters": len(text)
                }
            }
            
            return {
                "success": True,
                "content": content
            }
            
        except Exception as e:
            logger.error(f"Error extracting from text file {file_path}: {e}")
            return {
                "success": False,
                "error": f"Text file extraction error: {str(e)}"
            }
    
    def _extract_from_json(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from JSON files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Convert JSON to text representation
            text = json.dumps(data, indent=2)
            
            content = {
                "text": text,
                "metadata": {
                    "json_keys": list(data.keys()) if isinstance(data, dict) else [],
                    "data_type": type(data).__name__
                },
                "original_data": data
            }
            
            return {
                "success": True,
                "content": content
            }
            
        except Exception as e:
            logger.error(f"Error extracting from JSON file {file_path}: {e}")
            return {
                "success": False,
                "error": f"JSON file extraction error: {str(e)}"
            }
    
    def _structure_document(self, content: Dict[str, Any], file_path: Path, doc_type: str) -> Dict[str, Any]:
        """Structure the extracted content into a standardized format."""
        try:
            # Generate document ID
            doc_id = str(uuid.uuid4())
            
            # Create document hash for deduplication
            content_hash = hashlib.md5(content["text"].encode()).hexdigest()
            
            # Extract key information
            extracted_info = self._extract_key_information(content["text"], doc_type)
            
            # Create structured document
            structured_doc = {
                "id": doc_id,
                "title": extracted_info.get("title", file_path.stem),
                "content": content["text"],
                "source": str(file_path),
                "type": doc_type,
                "upload_date": datetime.now().isoformat(),
                "tags": extracted_info.get("tags", []),
                "metadata": {
                    "file_path": str(file_path),
                    "file_size": file_path.stat().st_size,
                    "file_extension": file_path.suffix,
                    "content_hash": content_hash,
                    "word_count": len(content["text"].split()),
                    "character_count": len(content["text"]),
                    "extracted_metadata": content.get("metadata", {}),
                    "key_topics": extracted_info.get("key_topics", []),
                    "entities": extracted_info.get("entities", []),
                    "summary": extracted_info.get("summary", "")
                }
            }
            
            return structured_doc
            
        except Exception as e:
            logger.error(f"Error structuring document: {e}")
            raise
    
    def _extract_key_information(self, text: str, doc_type: str) -> Dict[str, Any]:
        """Extract key information from document text."""
        try:
            # Basic text analysis
            sentences = text.split('.') if text else []
            words = text.split() if text else []
            
            # Extract potential title (first sentence or first line)
            title = sentences[0][:100] if sentences else "Untitled Document"
            
            # Generate tags based on document type and content
            tags = self._generate_tags(text, doc_type)
            
            # Extract key topics (simple keyword extraction)
            key_topics = self._extract_key_topics(text)
            
            # Extract entities (basic pattern matching)
            entities = self._extract_entities(text)
            
            # Generate summary (first few sentences)
            summary = '. '.join(sentences[:3]) if len(sentences) > 3 else text[:500]
            
            return {
                "title": title,
                "tags": tags,
                "key_topics": key_topics,
                "entities": entities,
                "summary": summary
            }
            
        except Exception as e:
            logger.error(f"Error extracting key information: {e}")
            return {
                "title": "Untitled Document",
                "tags": [],
                "key_topics": [],
                "entities": [],
                "summary": ""
            }
    
    def _generate_tags(self, text: str, doc_type: str) -> List[str]:
        """Generate tags based on document content and type."""
        tags = [doc_type]
        
        # Add domain-specific tags
        if doc_type == "safety":
            tags.extend(["safety", "hazard", "risk", "compliance"])
        elif doc_type == "engineering":
            tags.extend(["engineering", "design", "technical", "analysis"])
        elif doc_type == "regulatory":
            tags.extend(["regulatory", "compliance", "standards", "legal"])
        
        # Add content-based tags
        text_lower = text.lower()
        
        # Safety-related keywords
        safety_keywords = ["safety", "hazard", "risk", "danger", "toxic", "corrosive", "flammable"]
        for keyword in safety_keywords:
            if keyword in text_lower:
                tags.append(keyword)
        
        # Engineering keywords
        engineering_keywords = ["design", "analysis", "stress", "pressure", "temperature", "material"]
        for keyword in engineering_keywords:
            if keyword in text_lower:
                tags.append(keyword)
        
        # Remove duplicates and return
        return list(set(tags))
    
    def _extract_key_topics(self, text: str) -> List[str]:
        """Extract key topics from text using simple keyword analysis."""
        if not NLTK_AVAILABLE:
            return []
        
        try:
            # Tokenize and get parts of speech
            tokens = word_tokenize(text.lower())
            pos_tags = nltk.pos_tag(tokens)
            
            # Extract nouns and adjectives (potential topics)
            topics = []
            for word, pos in pos_tags:
                if pos.startswith('NN') or pos.startswith('JJ'):  # Nouns and adjectives
                    if len(word) > 3 and word.isalpha():  # Filter short words and non-alphabetic
                        topics.append(word)
            
            # Get most common topics
            from collections import Counter
            topic_counts = Counter(topics)
            return [topic for topic, count in topic_counts.most_common(10)]
            
        except Exception as e:
            logger.error(f"Error extracting key topics: {e}")
            return []
    
    def _extract_entities(self, text: str) -> List[str]:
        """Extract entities from text using pattern matching."""
        entities = []
        
        # Extract chemical formulas (e.g., H2SO4, CH3COOH)
        chemical_pattern = r'\b[A-Z][a-z]?\d*[A-Z][a-z]?\d*\b'
        chemicals = re.findall(chemical_pattern, text)
        entities.extend(chemicals)
        
        # Extract measurements (e.g., 100°C, 2.5 bar, 50L)
        measurement_pattern = r'\d+\.?\d*\s*(°C|bar|L|kg|MPa|psi)'
        measurements = re.findall(measurement_pattern, text)
        entities.extend(measurements)
        
        # Extract standards (e.g., OSHA, EPA, ASME)
        standard_pattern = r'\b(OSHA|EPA|ASME|API|DOT|UN|ISO|IEC)\b'
        standards = re.findall(standard_pattern, text, re.IGNORECASE)
        entities.extend(standards)
        
        return list(set(entities))
    
    def batch_process_documents(self, file_paths: List[str], doc_type: str = "general") -> Dict[str, Any]:
        """Process multiple documents in batch."""
        results = {
            "success": True,
            "processed": [],
            "errors": [],
            "total_files": len(file_paths),
            "successful": 0,
            "failed": 0
        }
        
        for file_path in file_paths:
            try:
                result = self.process_document(file_path, doc_type)
                if result["success"]:
                    results["processed"].append(result["document"])
                    results["successful"] += 1
                else:
                    results["errors"].append({
                        "file": file_path,
                        "error": result["error"]
                    })
                    results["failed"] += 1
            except Exception as e:
                results["errors"].append({
                    "file": file_path,
                    "error": str(e)
                })
                results["failed"] += 1
        
        if results["failed"] > 0:
            results["success"] = False
        
        return results  